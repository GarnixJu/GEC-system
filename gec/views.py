from django.http import JsonResponse
import spacy
import json
import os

from .translate import translate
from errant.parallel_to_diff import parallel_to_diff

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from django.utils.encoding import smart_str

nlp = spacy.load(os.environ.get('SPACY_MODEL', 'en'), disable=['ner'])


def correct_it(request):
    if request.is_ajax() and request.method == 'POST':
        data = json.loads(request.body.decode('utf-8'))
        text = data.get('text', '')
    elif request.method == 'GET':
        text = request.GET.get('text', '')
        data = {'text': text}
    else:
        data = {}

    lines = [' '.join(token.text for token in sent)
             for line in text.splitlines() if line.strip()
             for sent in nlp(line.strip()).sents
             ]
    tokenized_text = '\n'.join(lines)
    corrected_text = translate(tokenized_text)

    data['result'] = corrected_text
    diff = [parallel_to_diff(before, after, nlp)
            for before, after in zip(lines, corrected_text.splitlines())]
    data['word_diff'] = '\n'.join(diff)
    data['word_diff_by_sent'] = diff
    return JsonResponse(data)

def word_doc_correct_it(request):
    
    if request.is_ajax() and request.method == 'POST':
        data = json.loads(request.body.decode('utf-8'))
        text = data.get('text', '')
    elif request.method == 'GET':
        text = request.GET.get('text', '')
        data = {'text': text}
    else:
        data = {}

    lines = [' '.join(token.text for token in sent)
             for line in text.splitlines() if line.strip()
             for sent in nlp(line.strip()).sents
             ]
    tokenized_text = '\n'.join(lines)
    corrected_text = translate(tokenized_text)

    data['result'] = corrected_text
    diff = [parallel_to_diff(before, after, nlp)
            for before, after in zip(lines, corrected_text.splitlines())]
    data['word_diff'] = '\n'.join(diff)
    data['word_diff_by_sent'] = diff
    
    document = Document()

    document.add_heading('GEC Result', 0)

    p = document.add_paragraph('')
    
    print(data['word_diff_by_sent'])
    for line in data['word_diff_by_sent']:
        for token in line.split(' '):
            if token.startswith('{+'):
                after_word = token[2:-2].replace('\u3000',' ')
                p.add_run(after_word).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                p.add_run(' ')
            elif token.startswith('[-'):
                if token.endswith('+}'):
                    delete_tmp, insert_tmp = token[2:-2].split('-]{+')
                    delete = delete_tmp.replace('\u3000',' ')
                    insert = insert_tmp.replace('\u3000',' ')
                    p.add_run(insert).font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    p.add_run(delete).font.highlight_color = WD_COLOR_INDEX.RED
                    p.add_run(' ')
                else:
                    before_word = token[2:-2].replace('\u3000',' ')
                    p.add_run(before_word).font.highlight_color = WD_COLOR_INDEX.RED
                    p.add_run(' ')
            else:
                toke = token.replace('\u3000',' ')
                p.add_run(token)
                p.add_run(' ')


    document.add_page_break()

    document.save('gec.result.docx')
    file = open('./gec.result.docx', 'rb')
    response = HttpResponse(file)
    response['Content-Type'] = 'application/octet-stream' #設定頭資訊，告訴瀏覽器這是個檔案
    response['Content-Disposition'] = 'attachment;filename="gec.result.docx"'

    
    return response
